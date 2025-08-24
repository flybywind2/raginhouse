import os
import tempfile
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
from datetime import datetime
import requests
from fastapi import UploadFile
from src.clients.llm_client import LLMClient
from src.clients.document_client import DocumentClient
from src.models.schemas import ProcessedDocument, DocumentMetadata
from src.core.config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Document processing service using docling and LLM metadata generation

    초보자용 설명:
    - 업로드 파일/Confluence 페이지를 읽어 구조를 분석하고, 작은 청크로 나눕니다.
    - LLM으로 요약/토픽/언어 등 메타데이터를 생성합니다.
    - 최종 청크를 외부 문서 입력 API에 전송해 색인합니다.
    """
    
    def __init__(self):
        self.llm_client = LLMClient()
        self.document_client = DocumentClient()
        self.max_chunk_size = settings.MAX_CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
    
    async def process_uploaded_file(
        self,
        file: UploadFile,
        index_name: str
    ) -> bool:
        """Process uploaded file with docling and ingest to RAG system

        초보자용 설명:
        - 업로드된 파일을 임시로 저장한 후, docling으로 구조를 분석하고 청크로 분할하여 색인합니다.
        """
        temp_file_path = None
        
        try:
            # Save uploaded file temporarily
            suffix = Path(file.filename).suffix if file.filename else '.tmp'
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
                content = await file.read()
                temp_file.write(content)
                temp_file_path = temp_file.name
            
            doc_id = self._generate_doc_id(file.filename or "unknown")
            
            # Process document
            processed_doc = await self._process_document_file(
                temp_file_path,
                doc_id,
                file.filename or "unknown"
            )
            
            # Ingest to RAG system
            success = await self.document_client.insert_documents(
                processed_doc.chunks,
                index_name
            )
            
            if success:
                logger.info(f"Successfully processed and ingested file: {file.filename}")
            else:
                logger.error(f"Failed to ingest processed file: {file.filename}")
                
            return success
            
        except Exception as e:
            logger.error(f"File processing failed for {file.filename}: {e}")
            return False
        finally:
            # Clean up temporary file
            if temp_file_path and os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
    
    async def process_confluence_page(
        self,
        base_url: str,
        page_id: str,
        index_name: str
    ) -> bool:
        """Process Confluence page with single-page collection

        초보자용 설명:
        - Confluence API에서 페이지를 불러와 텍스트만 추출하고, 청크로 나누어 색인합니다.
        """
        try:
            # Fetch page content from Confluence
            page_data = await self._fetch_confluence_page(base_url, page_id)
            if not page_data:
                return False
            
            doc_id = f"confluence_{page_id}"
            
            # Process the page content
            processed_doc = await self._process_confluence_content(
                page_data,
                doc_id
            )
            
            # Ingest to RAG system
            success = await self.document_client.insert_documents(
                processed_doc.chunks,
                index_name
            )
            
            if success:
                logger.info(f"Successfully processed and ingested Confluence page: {page_id}")
            else:
                logger.error(f"Failed to ingest Confluence page: {page_id}")
                
            return success
            
        except Exception as e:
            logger.error(f"Confluence processing failed for page {page_id}: {e}")
            return False
    
    async def _process_document_file(
        self,
        file_path: str,
        doc_id: str,
        filename: str
    ) -> ProcessedDocument:
        """Process document file using docling"""
        
        return await self._process_with_docling(file_path, doc_id, filename)
    
    async def _process_with_docling(
        self,
        file_path: str,
        doc_id: str,
        filename: str
    ) -> ProcessedDocument:
        """Process document with docling for structure awareness

        초보자용 설명:
        - docling을 이용해 섹션/표/그림 정보를 인식하고, 섹션별 텍스트를 추출합니다.
        - 일부 텍스트 샘플을 LLM에 전달해 요약/토픽 등 메타데이터를 만듭니다.
        """
        try:
            # Import docling dynamically to handle optional dependency
            from docling.document_converter import DocumentConverter
            
            converter = DocumentConverter()
            result = converter.convert(file_path)
            doc = result.document
            
            # Extract structured content
            sections = []
            for section in doc.sections:
                section_data = {
                    'title': section.title,
                    'content': section.text,
                    'level': section.level,
                    'tables': [self._serialize_table(table) for table in section.tables],
                    'figures': [self._serialize_figure(fig) for fig in section.figures]
                }
                sections.append(section_data)
            
            # Generate metadata using LLM
            title = doc.title or filename
            content_sample = " ".join([s['content'][:500] for s in sections[:3]])  # First 3 sections
            
            metadata_dict = await self.llm_client.generate_metadata(
                title=title,
                content_sample=content_sample,
                doc_type=self._detect_doc_type(filename)
            )
            
            metadata = DocumentMetadata(
                doc_id=doc_id,
                title=title,
                summary=metadata_dict.get('summary', ''),
                topics=metadata_dict.get('topics', []),
                language=metadata_dict.get('language', 'unknown'),
                has_pii=metadata_dict.get('has_pii', False),
                doc_type=metadata_dict.get('doc_type', 'unknown'),
                created_time=datetime.now().isoformat(),
                permission_groups=[]
            )
            
            # Structure-aware chunking
            chunks = await self._chunk_with_structure_awareness(sections, metadata)
            
            return ProcessedDocument(metadata=metadata, chunks=chunks)
            
        except ImportError as e:
            logger.error("Docling not available")
            raise e
        except Exception as e:
            logger.error(f"Docling processing failed: {e}")
            raise e
    
    
    async def _fetch_confluence_page(self, base_url: str, page_id: str) -> Optional[Dict[str, Any]]:
        """Fetch Confluence page content

        초보자용 설명:
        - Confluence REST API를 호출해 페이지의 본문/메타데이터를 가져옵니다.
        """
        try:
            # Construct API URL
            api_url = f"{base_url.rstrip('/')}/rest/api/content/{page_id}?expand=body.storage,metadata.labels"
            
            headers = {}
            if settings.CONFLUENCE_TOKEN:
                headers["Authorization"] = f"Bearer {settings.CONFLUENCE_TOKEN}"
            
            # Make request with SSL verification disabled (MVP setting)
            response = requests.get(
                api_url,
                headers=headers,
                verify=settings.CONFLUENCE_VERIFY_SSL,
                timeout=30
            )
            
            if response.status_code == 200:
                return response.json()
            else:
                logger.error(f"Confluence API error: {response.status_code} - {response.text}")
                return None
                
        except Exception as e:
            logger.error(f"Confluence page fetch failed: {e}")
            return None
    
    async def _process_confluence_content(
        self,
        page_data: Dict[str, Any],
        doc_id: str
    ) -> ProcessedDocument:
        """Process Confluence page data

        초보자용 설명:
        - HTML 태그를 제거해 순수 텍스트를 만들고, 간단한 청크 나누기를 수행합니다.
        - LLM으로 메타데이터를 생성합니다.
        """
        try:
            title = page_data.get('title', 'Unknown Page')
            content = page_data.get('body', {}).get('storage', {}).get('value', '')
            
            # Simple HTML content extraction (could be enhanced)
            import re
            content = re.sub(r'<[^>]+>', '', content)  # Remove HTML tags
            content = re.sub(r'\s+', ' ', content).strip()  # Normalize whitespace
            
            # Generate metadata
            metadata_dict = await self.llm_client.generate_metadata(
                title=title,
                content_sample=content[:1000],
                doc_type='confluence_page'
            )
            
            metadata = DocumentMetadata(
                doc_id=doc_id,
                title=title,
                summary=metadata_dict.get('summary', ''),
                topics=metadata_dict.get('topics', []),
                language=metadata_dict.get('language', 'unknown'),
                has_pii=metadata_dict.get('has_pii', False),
                doc_type='confluence_page',
                created_time=datetime.now().isoformat(),
                permission_groups=['user']  # Default permission
            )
            
            # Simple chunking
            chunks = await self._simple_chunk(content, metadata)
            
            return ProcessedDocument(metadata=metadata, chunks=chunks)
            
        except Exception as e:
            logger.error(f"Confluence content processing failed: {e}")
            raise
    
    async def _chunk_with_structure_awareness(
        self,
        sections: List[Dict[str, Any]],
        metadata: DocumentMetadata
    ) -> List[Dict[str, Any]]:
        """Structure-aware chunking using docling sections

        초보자용 설명:
        - 섹션 단위로 묶되, 너무 길면 설정된 크기 기준으로 나눕니다.
        - 섹션 제목을 보존해 검색/인용 시 유용하게 씁니다.
        """
        chunks = []
        chunk_index = 0
        
        for section in sections:
            section_title = section.get('title', '')
            section_content = section.get('content', '')
            
            if not section_content.strip():
                continue
            
            # If section is small enough, keep as single chunk
            if len(section_content) <= self.max_chunk_size:
                chunk_doc = self._create_chunk_document(
                    metadata, section_content, chunk_index, section_title
                )
                chunks.append(chunk_doc)
                chunk_index += 1
            else:
                # Split large sections
                section_chunks = self._split_text(section_content)
                for chunk_text in section_chunks:
                    chunk_doc = self._create_chunk_document(
                        metadata, chunk_text, chunk_index, section_title
                    )
                    chunks.append(chunk_doc)
                    chunk_index += 1
        
        return chunks
    
    async def _simple_chunk(
        self,
        content: str,
        metadata: DocumentMetadata
    ) -> List[Dict[str, Any]]:
        """Simple text chunking

        초보자용 설명:
        - 길이에 따라 일정 크기로 텍스트를 자르고, 겹침(overlap)을 적용해 문맥 손실을 줄입니다.
        """
        chunks = []
        text_chunks = self._split_text(content)
        
        for i, chunk_text in enumerate(text_chunks):
            chunk_doc = self._create_chunk_document(metadata, chunk_text, i)
            chunks.append(chunk_doc)
        
        return chunks
    
    def _create_chunk_document(
        self,
        metadata: DocumentMetadata,
        content: str,
        chunk_index: int,
        section_title: Optional[str] = None
    ) -> Dict[str, Any]:
        """Create chunk document for insertion

        초보자용 설명:
        - 검색 API 요구사항에 맞춘 필드로 하나의 청크를 문서 형태로 만듭니다.
        """
        # Serialize metadata to flat string (required by API)
        additional_field = self._serialize_metadata(metadata, section_title)
        
        return {
            'doc_id': f"{metadata.doc_id}#{chunk_index:03d}",
            'title': metadata.title,
            'content': content.strip(),
            'section_title': section_title or '',
            'chunk_index': chunk_index,
            'permission_groups': metadata.permission_groups,
            'created_time': metadata.created_time,
            'additional_field': additional_field
        }
    
    def _serialize_metadata(
        self,
        metadata: DocumentMetadata,
        section_title: Optional[str] = None
    ) -> str:
        """Serialize metadata to flat string format (required by API)

        초보자용 설명:
        - 메타데이터를 단일 문자열로 펼쳐 저장합니다(외부 API 요구 형식).
        """
        parts = [
            f"summary={metadata.summary or ''}",
            f"topics={','.join(metadata.topics)}",
            f"language={metadata.language}",
            f"pii={str(metadata.has_pii).lower()}",
            f"section={section_title or ''}",
            f"doc_type={metadata.doc_type}"
        ]
        return " | ".join(parts)
    
    def _split_text(self, text: str) -> List[str]:
        """Split text into chunks with overlap

        초보자용 설명:
        - 문장 경계나 줄바꿈을 우선 고려해 자연스러운 분할을 시도합니다.
        - 다음 청크의 시작을 약간 겹치게 해 문맥 연속성을 높입니다.
        """
        if len(text) <= self.max_chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.max_chunk_size
            
            if end >= len(text):
                chunks.append(text[start:])
                break
            
            # Try to break at sentence boundary
            chunk = text[start:end]
            last_period = chunk.rfind('.')
            last_newline = chunk.rfind('\n')
            break_point = max(last_period, last_newline)
            
            if break_point > start + self.max_chunk_size // 2:
                end = start + break_point + 1
            
            chunks.append(text[start:end].strip())
            start = end - self.chunk_overlap
        
        return [chunk for chunk in chunks if chunk.strip()]
    
    def _serialize_table(self, table) -> str:
        """Serialize table to text representation

        초보자용 설명:
        - 간단히 캡션만 문자열로 보존합니다(자세한 구조화는 생략).
        """
        # Simplified table serialization
        return f"Table: {getattr(table, 'caption', 'No caption')}"
    
    def _serialize_figure(self, figure) -> str:
        """Serialize figure to text representation

        초보자용 설명:
        - 간단히 캡션만 문자열로 보존합니다.
        """
        # Simplified figure serialization
        return f"Figure: {getattr(figure, 'caption', 'No caption')}"
    
    def _detect_doc_type(self, filename: str) -> str:
        """Detect document type from filename

        초보자용 설명:
        - 파일 확장자에 따라 문서 유형을 추정합니다.
        """
        ext = Path(filename).suffix.lower()
        type_mapping = {
            '.pdf': 'pdf',
            '.docx': 'document',
            '.doc': 'document',
            '.pptx': 'presentation',
            '.ppt': 'presentation',
            '.xlsx': 'spreadsheet',
            '.xls': 'spreadsheet',
            '.txt': 'text',
            '.md': 'markdown'
        }
        return type_mapping.get(ext, 'unknown')
    
    def _generate_doc_id(self, filename: str) -> str:
        """Generate unique document ID

        초보자용 설명:
        - 파일명과 현재 시간을 합쳐 MD5 해시를 만들고, 앞부분만 ID로 씁니다.
        """
        import hashlib
        timestamp = datetime.now().isoformat()
        content = f"{filename}_{timestamp}"
        return hashlib.md5(content.encode()).hexdigest()[:12]
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF (simplified implementation)

        초보자용 설명:
        - PyPDF2로 페이지별 텍스트를 간단히 추출합니다(한계가 있습니다).
        """
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                return text
        except ImportError:
            logger.warning("PyPDF2 not available for PDF processing")
            return ""
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            return ""
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX (simplified implementation)

        초보자용 설명:
        - python-docx로 단락 텍스트를 읽어와 합칩니다.
        """
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except ImportError:
            logger.warning("python-docx not available for DOCX processing")
            return ""
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            return ""
